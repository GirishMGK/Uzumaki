"""
Word (.docx) Redactor — uses python-docx.

Strategy: for each paragraph, build the full text from all runs,
apply regex substitutions, then write the result back into the first run
(clearing the rest). This handles text that spans multiple runs.
Header/footer and table cells are also processed.
"""

import re
from pathlib import Path


def redact_word(
    input_path: str,
    output_path: str,
    patterns: list,
    custom_terms: list,
    replacement: str = "[REDACTED]",
    case_sensitive: bool = False,
) -> tuple:
    """Returns (total_count, audit_entries)."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for Word redaction.\n"
            "Install with:  pip install python-docx"
        )

    flags = 0 if case_sensitive else re.IGNORECASE

    # Pre-compile all patterns once
    compiled = []
    for regex, label in patterns:
        try:
            compiled.append((re.compile(regex, flags), label))
        except re.error:
            pass
    for kw in custom_terms:
        kw = kw.strip()
        if kw:
            try:
                compiled.append((re.compile(re.escape(kw), flags), "Custom Keyword"))
            except re.error:
                pass

    doc = Document(input_path)
    total = 0
    audit = []
    file_name = Path(input_path).name

    def _process_paragraph(para, location: str) -> int:
        if not para.runs:
            return 0
        full_text = "".join(r.text for r in para.runs)
        new_text = full_text
        count = 0
        for pat, label in compiled:
            matches = list(pat.finditer(new_text))
            if matches:
                for m in matches:
                    audit.append({
                        "File": file_name,
                        "Page": location,
                        "Category": label,
                        "Original Term": (m.group()[:45] + "…") if len(m.group()) > 45 else m.group(),
                        "Count": 1,
                    })
                new_text, n = pat.subn(replacement, new_text)
                count += n
        if count:
            # Put redacted text in first run, blank the rest
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        return count

    # ── Body paragraphs ───────────────────────────────────────────────────────
    for para in doc.paragraphs:
        total += _process_paragraph(para, "Body")

    # ── Tables ────────────────────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += _process_paragraph(para, "Table")

    # ── Headers & Footers ─────────────────────────────────────────────────────
    for section in doc.sections:
        for para in section.header.paragraphs:
            total += _process_paragraph(para, "Header")
        for para in section.footer.paragraphs:
            total += _process_paragraph(para, "Footer")

    doc.save(output_path)
    return total, audit
